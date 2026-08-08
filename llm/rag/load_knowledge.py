import json

from docx import Document
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import Table
from docx.text.paragraph import Paragraph

from config import get_resume_filename
from paths import ROOT


def load_knowledge_files() -> list:
    path = ROOT / "data" / "knowledge"

    knowledge: list = []

    for file in sorted(path.glob("*.json")):
        with file.open(encoding="utf-8") as f:
            knowledge.append(json.load(f))

    return knowledge


def iter_resume_paragraphs(doc: Document):
    """Yield paragraphs in document order, including those inside tables.

    Enhancv-style layouts often put all resume text in a 2-column table, so
    ``doc.paragraphs`` alone can be empty or nearly empty.
    """
    body = doc.element.body
    for child in body.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, doc)
        elif isinstance(child, CT_Tbl):
            table = Table(child, doc)
            seen_cells: set[int] = set()
            for row in table.rows:
                for cell in row.cells:
                    cell_id = id(cell._tc)
                    if cell_id in seen_cells:
                        continue
                    seen_cells.add(cell_id)
                    for paragraph in cell.paragraphs:
                        yield paragraph


def extract_resume_blocks() -> list:
    path = ROOT / "data" / "resume"
    file = get_resume_filename()
    blocks = []

    doc = Document(f"{path}/{file}")

    for index, paragraph in enumerate(iter_resume_paragraphs(doc)):
        text = paragraph.text.strip()

        if not text:
            continue

        blocks.append({
            "id": f"paragraph-{index}",
            "text": text,
            "style": paragraph.style.name if paragraph.style else "",
        })

    return blocks
