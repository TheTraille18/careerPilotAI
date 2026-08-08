from __future__ import annotations

import json
import re
from io import BytesIO

from docx import Document
from langchain_chroma import Chroma
from langchain_core.messages import HumanMessage, SystemMessage

from api.schemas import Job
from config import get_chroma_collection_name, get_chroma_persist_dir
from llm.bedrock.bedrock import llm_service
from llm.prompt.prompt import prompt_cover_letter
from llm.rag.embedding import create_embeddings
from llm.rag.load_knowledge import extract_resume_blocks
from s3_store import get_bucket_name, get_job_description, upload_tailored_resume


def _filename_part(value: str, fallback: str) -> str:
    text = re.sub(r"\s+", "_", (value or "").strip())
    text = re.sub(r"[^\w.\-]+", "", text)
    text = re.sub(r"_+", "_", text).strip("._-")
    return text or fallback


def build_cover_letter_filename(*, company: str, title: str) -> str:
    """Build Cover_Letter_<Company>_<Position>.docx with spaces as underscores."""
    company_part = _filename_part(company, "Company")
    title_part = _filename_part(title, "Position")
    return f"Cover_Letter_{company_part}_{title_part}.docx"


def parse_cover_letter_response(response_text: str) -> dict:
    text = (response_text or "").strip()
    if text.startswith("```"):
        text = text.split("\n", 1)[1]
        if "```" in text:
            text = text[: text.rfind("```")].strip()

    start = text.find("{")
    end = text.rfind("}")
    if start == -1 or end == -1 or end <= start:
        raise ValueError(f"No JSON object found in cover letter response: {text[:300]!r}")

    text = text[start : end + 1]
    try:
        payload = json.loads(text)
    except json.JSONDecodeError as exc:
        raise ValueError(
            "Cover letter model returned invalid/truncated JSON. "
            f"Parse error: {exc}. Snippet: {text[max(0, exc.pos - 40) : exc.pos + 40]!r}"
        ) from exc

    if not isinstance(payload, dict):
        raise ValueError("Cover letter response JSON must be an object")
    return payload


def cover_letter_to_docx(payload: dict) -> bytes:
    doc = Document()
    greeting = str(payload.get("greeting") or "Dear Hiring Manager,").strip()
    paragraphs = payload.get("bodyParagraphs") or []
    closing = str(payload.get("closing") or "Sincerely,").strip()
    signature = str(payload.get("signatureName") or "Justin Traille").strip()

    doc.add_paragraph(greeting)
    doc.add_paragraph("")

    if isinstance(paragraphs, list):
        for paragraph in paragraphs:
            text = str(paragraph or "").strip()
            if text:
                doc.add_paragraph(text)
    elif isinstance(paragraphs, str) and paragraphs.strip():
        doc.add_paragraph(paragraphs.strip())

    doc.add_paragraph("")
    doc.add_paragraph(closing)
    doc.add_paragraph(signature)

    buffer = BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def generate_cover_letter(job: Job) -> dict:
    full_job_description = get_job_description(job.jobId)
    resume_blocks = extract_resume_blocks()

    embedding = create_embeddings()
    vector_store = Chroma(
        collection_name=get_chroma_collection_name(),
        embedding_function=embedding,
        persist_directory=get_chroma_persist_dir(),
    )
    relevant_docs = vector_store.similarity_search(full_job_description)

    system_prompt, human_prompt = prompt_cover_letter(
        full_job_description,
        resume_blocks,
        relevant_docs,
        company=job.company or "",
        title=job.title or "",
    )

    llm = llm_service()
    response = llm.invoke(
        [
            SystemMessage(content=system_prompt),
            HumanMessage(content=human_prompt),
        ]
    )
    payload = parse_cover_letter_response(response.content)
    body = cover_letter_to_docx(payload)
    filename = build_cover_letter_filename(
        company=job.company or "",
        title=job.title or "",
    )

    # Same S3 directory as tailored resumes: jobs/<job_id>/resume/
    s3_key = upload_tailored_resume(
        job_id=job.jobId,
        body=body,
        filename=filename,
    )

    return {
        "bucket": get_bucket_name(),
        "key": s3_key,
        "filename": filename,
    }
